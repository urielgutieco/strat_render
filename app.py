import os
import io
import zipfile
import random
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from flask_cors import CORS
from flask_mail import Mail, Message
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
CORS(app)

# --- CONFIGURACIÓN DE PRODUCCIÓN ---
app.config['SECRET_KEY'] = os.environ.get("SESSION_SECRET", "una_clave_muy_segura_123")

# --- CONFIGURACIÓN DE BASE DE DATOS ---
db_url = os.environ.get("DATABASE_URL", "sqlite:///usuarios.db")
if db_url.startswith("postgres://"):
    # Render usa 'postgres://', pero SQLAlchemy requiere 'postgresql://'
    # Además, añadimos el parámetro de SSL para producción
    db_url = db_url.replace("postgres://", "postgresql://", 1)

app.config['SQLALCHEMY_DATABASE_URI'] = db_url
# Añadir esto para evitar desconexiones en Render:
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {"pool_pre_ping": True}

# --- CONFIGURACIÓN DE CORREO (Flask-Mail) ---
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.environ.get("MAIL_USERNAME")
app.config['MAIL_PASSWORD'] = os.environ.get("MAIL_PASSWORD")
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get("MAIL_USERNAME")

db = SQLAlchemy(app)
mail = Mail(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# --- CONFIGURACIÓN DE DIRECTORIOS ---
TEMPLATE_FOLDER = 'template_word'
TMP_DIR = '/tmp'
for folder in [TEMPLATE_FOLDER, TMP_DIR]:
    os.makedirs(folder, exist_ok=True)

# --- MODELO DE USUARIO ---
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    is_active = db.Column(db.Boolean, default=True) 
    is_admin = db.Column(db.Boolean, default=False)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# --- MAPEO DE SERVICIOS ---
SERVICIO_TO_DIR = {
    "Servicios de construccion de unidades unifamiliares": "construccion_unifamiliar",
    "Servicios de reparacion o ampliacion o remodelacion de viviendas unifamiliares": "reparacion_remodelacion_unifamiliar",
    "Servicio de remodelacion general de viviendas unifamiliares": "remodelacion_general",
    "Servicios de reparacion de casas moviles en el sitio": "reparacion_casas_moviles",
    "Servicios de construccion y reparacion de patios y terrazas": "patios_terrazas",
    "Servico de reparacion por daños ocasionados por fuego de viviendas unifamiliares": "reparacion_por_fuego",
    "Servicio de construccion de casas unifamiliares nuevas": "construccion_unifamiliar_nueva",
    "Servicio de instalacion de casas unifamiliares prefabricadas": "instalacion_prefabricadas",
    "Servicio de construccion de casas en la ciudad o casas jardin unifamiliares nuevas": "construccion_casas_ciudad_jardin",
    "Dasarrollo urbano": "desarrollo_urbano",
    "Servicio de planificacion de la ordenacion urbana": "planificacion_ordenacion_urbana",
    "Servicio de administracion de tierras urbanas": "administracion_tierras_urbanas",
    "Servicio de programacion de inversiones urbanas": "programacion_inversiones_urbanas",
    "Servicio de reestructuracion de barrios marginales": "reestructuracion_barrios_marginales",
    "Servicios de alumbrado urbano": "alumbrado_urbano",
    "Servicios de control o regulacion del desarrollo urbano": "control_desarrollo_urbano",
    "Servicios de estandares o regulacion de edificios urbanos": "estandares_regulacion_edificios",
    "Servicios comunitarios urbanos": "comunitarios_urbanos",
    "Servicios de administracion o gestion de proyectos o programas urbanos": "gestion_proyectos_programas_urbanos",
    "Ingenieria civil": "ingenieria_civil",
    "Ingenieria de carreteras": "ingenieria_carreteras",
    "Ingenieria deinfraestructura de instalaciones o fabricas": "infraestructura_instalaciones_fabricas",
    "Servicios de mantenimiento e instalacion de equipo pesado": "mantenimiento_instalacion_equipo_pesado",
    "Servicio de mantenimiento y reparacion de equipo pesado": "mantenimiento_reparacion_equipo_pesado",
}

TEMPLATE_FILES = ['plantilla_solicitud.docx', '2.docx', '3.docx', '4.docx', '1.docx']

# --- LÓGICA DE PROCESAMIENTO WORD ---
def replace_text_in_document(document, replacements):
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))

def generate_single_document(template_filename, template_root, replacements, user_image_path=None, data=None):
    template_path = os.path.join(template_root, template_filename)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Plantilla '{template_filename}' no encontrada.")

    document = Document(template_path)
    replace_text_in_document(document, replacements)

    if user_image_path and os.path.exists(user_image_path):
        try:
            document.add_paragraph()
            document.add_paragraph(data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A'))
            document.add_picture(user_image_path, width=Inches(2.5))
        except Exception:
            document.add_paragraph("⚠ Firma no disponible.")
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- RUTAS ---

@app.route('/')
def home():
    # Esto busca 'templates/index.html'
    return render_template('index.html')

@app.route('/login')
def login():
    # Esto busca 'templates/login.html'
    return render_template('login.html')

@app.route('/')
@login_required
def formulario():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('formulario'))
        
    if request.method == 'POST':
        user = User.query.filter_by(username=request.form.get('user')).first()
        password = request.form.get('password')
        
        if user and check_password_hash(user.password, password):
            if not user.is_active:
                flash("Membresía inactiva. Contacta al administrador.")
                return redirect(url_for('login'))
            login_user(user)
            return redirect(url_for('formulario'))
        
        flash("Usuario o contraseña incorrectos")
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/generate-word', methods=['POST'])
@login_required
def generate_word():
    try:
        data = request.form.to_dict()
        uploaded_image = request.files.get("imagen_usuario")
        user_image_path = None
        
        if uploaded_image and uploaded_image.filename != '':
            user_image_path = os.path.join(TMP_DIR, f"firma_{current_user.id}.png")
            uploaded_image.save(user_image_path)

        servicio = data.get('servicio')
        carpeta_servicio = SERVICIO_TO_DIR.get(servicio)
        
        if not carpeta_servicio:
            return jsonify({"error": "Servicio no válido"}), 400

        template_root = os.path.join(TEMPLATE_FOLDER, carpeta_servicio)
        numero_contrato = ''.join([str(random.randint(0, 9)) for _ in range(18)])
        rfc_cliente = data.get('r_f_c', 'N/A')

        replacements = {
            '${descripcion_del_servicio}': servicio,
            '${razon_social}': data.get('razon_social', 'N/A'),
            '${r_f_c}': rfc_cliente,
            '${domicilio_del_cliente}': data.get('domicilio_del_cliente', 'N/A'),
            '${telefono_del__cliente}': data.get('telefono_del__cliente', 'N/A'),
            '${correo_del_cliente}': data.get('correo_del_cliente', 'N/A'),
            '${fecha_de_inicio_del_servicio}': data.get('fecha_de_inicio_del_servicio', 'N/A'),
            '${fecha_de_conclusion_del_servicio}': data.get('fecha_de_conclusion_del_servicio', 'N/A'),
            '${monto_de_la_operacion_Sin_IVA}': data.get('monto_de_la_operacion_Sin_IVA', 'N/A'),
            '${forma_de_pago}': data.get('forma_de_pago', 'N/A'),
            '${cantidad}': data.get('cantidad', 'N/A'),
            '${unidad}': data.get('unidad', 'N/A'),
            '${numero_de_contrato}': numero_contrato,
            '${fecha_de_operación}': data.get('fecha_de_operación', 'N/A'),
            '${nombre_completo_de_la_persona_que_firma_la_solicitud}': data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${cargo_de_la_persona_que_firma_la_solicitud}': data.get('cargo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${factura_relacionada_con_la_operación}': data.get('factura_relacionada_con_la_operación', 'N/A'),
            '${informe_si_cuenta_con_fotografias_videos_o_informacion_adicion}': data.get('informe_si_cuenta_con_fotografias_videos_o_informacion_adicion', 'N/A'),
            '${comentarios}': data.get('comentarios', 'N/A')
        }

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for template in TEMPLATE_FILES:
                try:
                    doc_buf = generate_single_document(template, template_root, replacements, user_image_path, data)
                    base_name = os.path.splitext(template)[0]
                    filename = f"{base_name}_{numero_contrato}_{rfc_cliente}.docx"
                    zip_file.writestr(filename, doc_buf.getvalue())
                except Exception as e:
                    print(f"Error procesando {template}: {e}")

        zip_content = zip_buffer.getvalue()
        zip_filename = f"Contratos_{rfc_cliente}.zip"

        # --- ENVÍO DE CORREO ---
        # Definir los dos correos de destino aquí o recibirlos del form
        destinatarios = ["uriel.gutierrenz@gmail.com", "urielgutieco@gmail.com"] 
        
        try:
            msg = Message(
                subject=f"Nuevos Contratos Generados - RFC: {rfc_cliente}",
                recipients=destinatarios,
                body=f"Se adjuntan los documentos generados para el contrato {numero_contrato}."
            )
            msg.attach(zip_filename, "application/zip", zip_content)
            mail.send(msg)
            print("Correo enviado con éxito")
        except Exception as mail_err:
            print(f"Error enviando correo: {mail_err}")
            # Continuamos para que el usuario al menos pueda descargar el archivo

        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=zip_filename
        )

    except Exception as e:
        return jsonify({"error": f"Error interno: {str(e)}"}), 500

# --- INICIALIZACIÓN ---
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        if not User.query.filter_by(username="admin").first():
            admin = User(
                username="admin", 
                password=generate_password_hash("admin123"), 
                is_admin=True
            )
            db.session.add(admin)
            db.session.commit()
            print("Admin creado satisfactoriamente.")
            
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)